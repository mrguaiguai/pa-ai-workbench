import mimetypes
import re
from pathlib import Path
from typing import Any

from knowledge_engine.parsers.base import DocumentParser
from knowledge_engine.parsers.errors import DocumentParseError
from knowledge_engine.parsers.errors import ParserDependencyError
from knowledge_engine.parsers.errors import UnsupportedDocumentFormatError
from knowledge_engine.parsers.schemas import ParsedDocument
from knowledge_engine.parsers.schemas import ParsedSection


class FileDocumentParser(DocumentParser):
    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".txt": "txt",
        ".text": "txt",
        ".md": "markdown",
        ".markdown": "markdown",
    }
    MARKDOWN_HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")

    def parse(self, file_path: str, metadata: dict | None = None) -> ParsedDocument:
        path = Path(file_path)
        if not path.is_file():
            raise DocumentParseError(f"Document file does not exist: {path}")

        resolved_metadata = metadata or {}
        file_type = self._detect_file_type(path, resolved_metadata)
        if file_type == "pdf":
            return self._parse_pdf(path, resolved_metadata)
        if file_type == "docx":
            return self._parse_docx(path, resolved_metadata)
        if file_type == "markdown":
            return self._parse_markdown(path, resolved_metadata)
        if file_type == "txt":
            return self._parse_text(path, resolved_metadata)

        raise UnsupportedDocumentFormatError(
            f"Unsupported document format for {path.name}: {path.suffix or 'unknown'}"
        )

    def _parse_text(self, path: Path, metadata: dict) -> ParsedDocument:
        text, encoding = self._read_text(path)
        section = ParsedSection(
            text=text,
            title=metadata.get("title") or path.stem,
            start_char=0,
            end_char=len(text),
            metadata={"encoding": encoding},
        )
        return self._build_document(
            path=path,
            file_type="txt",
            text=text,
            sections=[section],
            metadata={**metadata, "encoding": encoding},
        )

    def _parse_markdown(self, path: Path, metadata: dict) -> ParsedDocument:
        text, encoding = self._read_text(path)
        sections = self._markdown_sections(text)
        return self._build_document(
            path=path,
            file_type="markdown",
            text=text,
            sections=sections,
            metadata={**metadata, "encoding": encoding},
        )

    def _parse_pdf(self, path: Path, metadata: dict) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ParserDependencyError(
                "pypdf is required to parse PDF documents. Install backend requirements."
            ) from exc

        try:
            reader = PdfReader(str(path))
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception as exc:  # pragma: no cover - pypdf varies here
                    raise DocumentParseError("Encrypted PDF could not be decrypted") from exc

            sections: list[ParsedSection] = []
            text_parts: list[str] = []
            cursor = 0
            for page_index, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                if page_text and not page_text.endswith("\n"):
                    page_text = f"{page_text}\n"
                start_char = cursor
                end_char = start_char + len(page_text)
                sections.append(
                    ParsedSection(
                        text=page_text,
                        title=f"Page {page_index}",
                        page_number=page_index,
                        start_char=start_char,
                        end_char=end_char,
                        metadata={"page_number": page_index},
                    )
                )
                text_parts.append(page_text)
                cursor = end_char
        except DocumentParseError:
            raise
        except Exception as exc:
            raise DocumentParseError(f"Failed to parse PDF document: {path.name}") from exc

        text = "".join(text_parts)
        return self._build_document(
            path=path,
            file_type="pdf",
            text=text,
            sections=sections,
            metadata={**metadata, "page_count": len(sections)},
        )

    def _parse_docx(self, path: Path, metadata: dict) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise ParserDependencyError(
                "python-docx is required to parse DOCX documents. Install backend requirements."
            ) from exc

        try:
            document = Document(str(path))
            blocks = self._docx_blocks(document)
        except Exception as exc:
            raise DocumentParseError(f"Failed to parse DOCX document: {path.name}") from exc

        text, sections = self._sections_from_blocks(blocks, default_title=path.stem)
        return self._build_document(
            path=path,
            file_type="docx",
            text=text,
            sections=sections,
            metadata={**metadata, "block_count": len(blocks)},
        )

    def _build_document(
        self,
        path: Path,
        file_type: str,
        text: str,
        sections: list[ParsedSection],
        metadata: dict,
    ) -> ParsedDocument:
        mime_type = metadata.get("mime_type") or mimetypes.guess_type(path.name)[0]
        return ParsedDocument(
            source_path=str(path),
            file_name=path.name,
            file_type=file_type,
            title=metadata.get("title") or self._title_from_sections(sections) or path.stem,
            text=text,
            sections=sections,
            mime_type=mime_type,
            metadata={
                **metadata,
                "parser": "file",
                "char_count": len(text),
                "section_count": len(sections),
            },
        )

    def _detect_file_type(self, path: Path, metadata: dict) -> str:
        explicit = str(metadata.get("file_type") or metadata.get("document_type") or "").lower()
        if explicit in {"pdf", "docx", "txt", "text", "markdown", "md"}:
            return "markdown" if explicit in {"markdown", "md"} else "txt" if explicit == "text" else explicit

        suffix = path.suffix.lower()
        if suffix in self.SUPPORTED_EXTENSIONS:
            return self.SUPPORTED_EXTENSIONS[suffix]

        mime_type = str(metadata.get("mime_type") or mimetypes.guess_type(path.name)[0] or "")
        if mime_type == "application/pdf":
            return "pdf"
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return "docx"
        if mime_type in {"text/plain", "text/markdown", "text/x-markdown"}:
            return "markdown" if "markdown" in mime_type else "txt"
        return "unsupported"

    @staticmethod
    def _read_text(path: Path) -> tuple[str, str]:
        raw = path.read_bytes()
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return raw.decode(encoding).replace("\r\n", "\n").replace("\r", "\n"), encoding
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="replace").replace("\r\n", "\n").replace("\r", "\n"), "utf-8-replace"

    def _markdown_sections(self, text: str) -> list[ParsedSection]:
        lines = text.splitlines(keepends=True)
        if not lines:
            return [ParsedSection(text="", start_char=0, end_char=0)]

        sections: list[ParsedSection] = []
        headings: list[str] = []
        current_lines: list[str] = []
        current_title: str | None = None
        current_path: list[str] = []
        current_start = 0
        cursor = 0
        in_fence = False

        def flush(end_char: int) -> None:
            nonlocal current_lines, current_title, current_path, current_start
            if not current_lines:
                return
            section_text = "".join(current_lines)
            sections.append(
                ParsedSection(
                    text=section_text,
                    title=current_title,
                    section_path=list(current_path),
                    start_char=current_start,
                    end_char=end_char,
                    metadata={"format": "markdown"},
                )
            )
            current_lines = []

        for line in lines:
            stripped = line.strip()
            if stripped.startswith("```") or stripped.startswith("~~~"):
                in_fence = not in_fence
            heading_match = None if in_fence else self.MARKDOWN_HEADING_PATTERN.match(stripped)
            if heading_match:
                flush(cursor)
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                headings = headings[: level - 1]
                headings.append(title)
                current_title = title
                current_path = list(headings)
                current_start = cursor
            elif not current_lines:
                current_start = cursor
            current_lines.append(line)
            cursor += len(line)

        flush(cursor)
        if sections:
            return sections
        return [
            ParsedSection(
                text=text,
                title=None,
                start_char=0,
                end_char=len(text),
                metadata={"format": "markdown"},
            )
        ]

    def _docx_blocks(self, document: Any) -> list[dict]:
        blocks: list[dict] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = getattr(paragraph.style, "name", "") or ""
            heading_level = self._heading_level(style_name)
            blocks.append(
                {
                    "text": text,
                    "title": text if heading_level is not None else None,
                    "heading_level": heading_level,
                    "style": style_name,
                }
            )

        for table_index, table in enumerate(document.tables, start=1):
            rows: list[str] = []
            for row in table.rows:
                cells = [" ".join(cell.text.split()) for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                blocks.append(
                    {
                        "text": "\n".join(rows),
                        "title": f"Table {table_index}",
                        "heading_level": None,
                        "style": "table",
                    }
                )
        return blocks

    def _sections_from_blocks(
        self,
        blocks: list[dict],
        default_title: str,
    ) -> tuple[str, list[ParsedSection]]:
        if not blocks:
            return "", [
                ParsedSection(
                    text="",
                    title=default_title,
                    start_char=0,
                    end_char=0,
                    metadata={"format": "docx"},
                )
            ]

        sections: list[ParsedSection] = []
        text_parts: list[str] = []
        current_parts: list[str] = []
        current_title: str | None = default_title
        current_path: list[str] = []
        current_start = 0
        cursor = 0
        headings: list[str] = []

        def append_part(part: str) -> int:
            nonlocal cursor
            if text_parts:
                text_parts.append("\n\n")
                cursor += 2
            part_start = cursor
            text_parts.append(part)
            cursor += len(part)
            return part_start

        def flush() -> None:
            nonlocal current_parts, current_title, current_path, current_start
            if not current_parts:
                return
            section_text = "\n\n".join(current_parts)
            sections.append(
                ParsedSection(
                    text=section_text,
                    title=current_title,
                    section_path=list(current_path),
                    start_char=current_start,
                    end_char=current_start + len(section_text),
                    metadata={"format": "docx"},
                )
            )
            current_parts = []

        for block in blocks:
            block_text = str(block["text"])
            heading_level = block.get("heading_level")
            if heading_level is not None:
                flush()
                headings = headings[: heading_level - 1]
                headings.append(block_text)
                current_title = block_text
                current_path = list(headings)
                current_start = append_part(block_text)
            elif not current_parts:
                current_start = append_part(block_text)
            else:
                append_part(block_text)
            current_parts.append(block_text)

        flush()
        return "".join(text_parts), sections

    @staticmethod
    def _heading_level(style_name: str) -> int | None:
        normalized = style_name.strip().lower()
        if not normalized.startswith("heading"):
            return None
        digits = "".join(char for char in normalized if char.isdigit())
        if not digits:
            return 1
        return int(digits)

    @staticmethod
    def _title_from_sections(sections: list[ParsedSection]) -> str | None:
        for section in sections:
            if section.title:
                return section.title
        return None
